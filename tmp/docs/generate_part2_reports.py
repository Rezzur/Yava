from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt


OUT_DIR = Path("C:/Yavimax3/output/doc")
OUT_DIR.mkdir(parents=True, exist_ok=True)

INSTITUTE_LINES = [
    "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
    "ФГБОУ ВО «МИРЭА - Российский технологический университет»",
    "Институт информационных технологий",
    "Кафедра математического обеспечения и стандартизации информационных технологий",
]


def set_run_font(run, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold


def set_paragraph_base(paragraph, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, indent=True):
    paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)


def add_paragraph(doc, text, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, bold=False, indent=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, bold=bold)
    set_paragraph_base(p, align=align, indent=indent)
    return p


def add_blank(doc):
    p = doc.add_paragraph()
    set_paragraph_base(p, align=WD_PARAGRAPH_ALIGNMENT.LEFT, indent=False)
    return p


def configure_document(doc):
    sec = doc.sections[0]
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    npf = normal.paragraph_format
    npf.line_spacing = 1.5
    npf.space_before = Pt(0)
    npf.space_after = Pt(0)
    npf.first_line_indent = Cm(1.25)
    npf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def add_title_page(doc, developer_name, role):
    for line in INSTITUTE_LINES:
        add_paragraph(doc, line, align=WD_PARAGRAPH_ALIGNMENT.CENTER, indent=False)

    add_blank(doc)
    add_blank(doc)

    add_paragraph(doc, "ОТЧЕТ", align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, indent=False)
    add_paragraph(
        doc,
        "по выполнению Части 2 итогового проекта",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        bold=True,
        indent=False,
    )
    add_paragraph(
        doc,
        "по дисциплине «Технологии разработки программных приложений»",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        indent=False,
    )

    add_blank(doc)
    add_paragraph(
        doc,
        "Тема проекта: «Кроссплатформенный десктоп-мессенджер Yavimax»",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        indent=False,
    )

    add_blank(doc)
    add_paragraph(doc, "Выполнил: студент группы [УКАЗАТЬ ГРУППУ]", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, indent=False)
    add_paragraph(doc, developer_name, align=WD_PARAGRAPH_ALIGNMENT.RIGHT, indent=False)
    add_paragraph(doc, f"Роль: {role}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, indent=False)
    add_blank(doc)
    add_paragraph(doc, "Проверил: [ФИО преподавателя]", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, indent=False)

    for _ in range(8):
        add_blank(doc)

    add_paragraph(doc, "Москва 2026", align=WD_PARAGRAPH_ALIGNMENT.CENTER, indent=False)
    doc.add_page_break()


def add_section_title(doc, title):
    add_paragraph(doc, title, align=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True, indent=False)


def style_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_paragraph_base(p, align=WD_PARAGRAPH_ALIGNMENT.CENTER, indent=False)
                if not p.runs:
                    p.add_run("")
                for run in p.runs:
                    set_run_font(run)


def add_commits_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"

    headers = ["Дата", "Хэш коммита", "Описание"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h

    for idx, row in enumerate(rows, start=1):
        table.cell(idx, 0).text = row[0]
        table.cell(idx, 1).text = row[1]
        table.cell(idx, 2).text = row[2]

    style_table(table)

    for i in range(1, 1 + len(rows)):
        p = table.cell(i, 2).paragraphs[0]
        set_paragraph_base(p, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, indent=False)


def add_figures(doc, figures):
    for fig in figures:
        add_paragraph(doc, fig["ref"])
        add_paragraph(doc, f"[ВСТАВИТЬ СКРИНШОТ: {fig['placeholder']}]", align=WD_PARAGRAPH_ALIGNMENT.LEFT, indent=False)
        add_paragraph(doc, fig["caption"], align=WD_PARAGRAPH_ALIGNMENT.CENTER, indent=False)


def add_timeline(doc, items):
    for i, item in enumerate(items, start=1):
        add_paragraph(doc, f"{i}. {item}")


def build_report(cfg):
    doc = Document()
    configure_document(doc)

    add_title_page(doc, cfg["developer_name"], cfg["role"])

    add_section_title(doc, "1. Введение")
    for p in cfg["intro"]:
        add_paragraph(doc, p)

    add_section_title(doc, "2. Основная часть")
    for i, p in enumerate(cfg["main"], start=1):
        add_paragraph(doc, f"{i}. {p}")

    add_section_title(doc, "3. Таблица коммитов")
    add_commits_table(doc, cfg["commits"])

    add_blank(doc)
    add_section_title(doc, "4. Иллюстративный материал")
    add_figures(doc, cfg["figures"])

    add_section_title(doc, "5. Краткая проделанная работа (с нуля до актуальной версии)")
    add_timeline(doc, cfg["timeline"])

    add_section_title(doc, "6. Заключение")
    for p in cfg["conclusion"]:
        add_paragraph(doc, p)

    out_path = OUT_DIR / cfg["filename"]
    doc.save(out_path)
    return out_path


reports = [
    {
        "filename": "Отчет_Часть2_Зырянов_ЯА.docx",
        "developer_name": "Зырянов Ян Александрович",
        "role": "Team Lead, Backend-разработчик",
        "intro": [
            "Целью выполнения Части 2 являлась организация инженерной базы backend-модуля проекта Yavimax: настройка системы сборки на Gradle, документирование серверного кода в формате Javadoc и формирование единого README для команды.",
            "Дополнительно была обеспечена воспроизводимая генерация HTML-документации для backend-части, что необходимо для проверки и сопровождения проекта.",
        ],
        "main": [
            "Настроена система сборки в директории backend на основе Gradle: создан и доработан build.gradle, подключены зависимости Spring Boot и WebSocket-модуля. Это позволило стандартизировать сборку и запуск сервера командой ./gradlew bootRun.",
            "Подготовлен и размещен README.md в корне репозитория с описанием архитектуры, зависимостей и командами запуска всех модулей проекта (backend, frontend, frontend-web).",
            "Настроена задача javadoc в backend/build.gradle с поддержкой кодировки UTF-8 для корректного отображения русскоязычных комментариев. Добавлены Javadoc-комментарии в конфигурационные классы WebSocket.",
            "Проверена генерация HTML-документации командой ./gradlew javadoc с формированием выходных файлов в каталоге backend/build/docs/javadoc.",
        ],
        "commits": [
            ["[ДД.ММ.2026]", "[HASH_1]", "Инициализация Gradle-сборки в backend, добавление Spring Boot и WebSocket-зависимостей"],
            ["[ДД.ММ.2026]", "[HASH_2]", "Создание README.md с архитектурой проекта и инструкцией запуска ./gradlew bootRun"],
            ["[ДД.ММ.2026]", "[HASH_3]", "Настройка javadoc (UTF-8) и добавление Javadoc-комментариев для WebSocket-конфигурации"],
        ],
        "figures": [
            {
                "ref": "Конфигурация backend-сборки представлена на рисунке 1.",
                "placeholder": "Файл backend/build.gradle: плагины, зависимости Spring Boot/WebSocket и блок javadoc",
                "caption": "Рисунок 1 - Конфигурация сборки backend-модуля в Gradle",
            },
            {
                "ref": "Пример документирования серверного кода приведен на рисунке 2.",
                "placeholder": "Фрагмент Java-кода с Javadoc-комментариями в WebSocket-конфигурации",
                "caption": "Рисунок 2 - Javadoc-комментарии в backend-классе",
            },
            {
                "ref": "Результат генерации документации в консоли показан на рисунке 3.",
                "placeholder": "Консоль с выполнением ./gradlew javadoc и строкой BUILD SUCCESSFUL",
                "caption": "Рисунок 3 - Успешная генерация Javadoc",
            },
            {
                "ref": "Итоговая HTML-документация представлена на рисунке 4.",
                "placeholder": "Открытый в браузере файл backend/build/docs/javadoc/index.html",
                "caption": "Рисунок 4 - Сгенерированная HTML-документация backend",
            },
            {
                "ref": "История коммитов разработчика представлена на рисунке 5.",
                "placeholder": "Вывод git log --author=\"Зырянов|Yan\" -n 3 --oneline",
                "caption": "Рисунок 5 - История коммитов разработчика 1",
            },
        ],
        "timeline": [
            "На стартовом этапе проект имел только базовую структуру репозитория и требовал формализации процесса сборки backend-части.",
            "Далее была настроена Gradle-конфигурация сервера, подключены необходимые зависимости Spring Boot и WebSocket для запуска базового функционала обмена сообщениями.",
            "После стабилизации сборки подготовлен единый README.md с инструкциями по развертыванию и запуску всех модулей команды.",
            "На следующем шаге внедрено Javadoc-документирование ключевых backend-компонентов и настроена корректная генерация HTML-документации в UTF-8.",
            "В актуальной версии backend-модуль имеет воспроизводимую сборку, документированный код и готовый артефакт технической документации для сопровождения.",
        ],
        "conclusion": [
            "По итогам Части 2 была сформирована полноценная инженерная база backend-модуля Yavimax: настроена сборка Gradle, подготовлен README и внедрено Javadoc-документирование с генерацией HTML.",
            "Поставленные задачи по роли Team Lead и Backend-разработчика выполнены в полном объеме.",
        ],
    },
    {
        "filename": "Отчет_Часть2_Топал_ВВ.docx",
        "developer_name": "Топал Виктор Викторович",
        "role": "Frontend UI/Electron-разработчик",
        "intro": [
            "Цель Части 2 в рамках роли Frontend UI/Electron-разработчика заключалась в создании воспроизводимой среды сборки клиентского модуля, документировании TypeScript-кода и настройке генерации HTML-документации.",
            "Отдельное внимание было уделено прозрачности запуска и сопровождения фронтенда через npm-скрипты и описание в README.",
        ],
        "main": [
            "В директории frontend инициализирована клиентская сборка: подготовлены package.json и tsconfig.json, добавлены зависимости для разработки интерфейса и Electron-окружения.",
            "В TypeScript-компонентах desktop-интерфейса добавлены TSDoc-комментарии с описанием назначения, параметров и возвращаемых значений.",
            "README.md дополнен инструкциями по запуску фронтенда и сопровождению модуля.",
            "В scripts добавлена команда npm run docs на основе TypeDoc для автоматической генерации HTML-документации.",
        ],
        "commits": [
            ["[ДД.ММ.2026]", "[HASH_1]", "Инициализация npm (package.json) и TypeScript (tsconfig.json) в frontend, добавление зависимостей Electron"],
            ["[ДД.ММ.2026]", "[HASH_2]", "Добавление TSDoc-комментариев в TypeScript-компоненты десктопного интерфейса"],
            ["[ДД.ММ.2026]", "[HASH_3]", "Дополнение README.md инструкциями по запуску фронтенда и добавление npm run docs (TypeDoc)"],
        ],
        "figures": [
            {
                "ref": "Файл сборки frontend-модуля представлен на рисунке 1.",
                "placeholder": "Файл frontend/package.json: scripts, зависимости и команда docs",
                "caption": "Рисунок 1 - Конфигурация npm-сборки frontend",
            },
            {
                "ref": "Пример TSDoc-комментариев показан на рисунке 2.",
                "placeholder": "Фрагмент TypeScript/TSX-кода с TSDoc-блоком комментариев",
                "caption": "Рисунок 2 - Документирование TypeScript-кода (TSDoc)",
            },
            {
                "ref": "Вывод генерации документации в консоли приведен на рисунке 3.",
                "placeholder": "Консоль с выполнением npm run docs и успешным завершением TypeDoc",
                "caption": "Рисунок 3 - Успешная генерация TypeDoc",
            },
            {
                "ref": "Итоговая HTML-документация представлена на рисунке 4.",
                "placeholder": "Открытый в браузере файл frontend/docs/index.html",
                "caption": "Рисунок 4 - Сгенерированная HTML-документация frontend",
            },
            {
                "ref": "История коммитов разработчика отражена на рисунке 5.",
                "placeholder": "Вывод git log --author=\"Топал|Viktor\" -n 3 --oneline",
                "caption": "Рисунок 5 - История коммитов разработчика 2",
            },
        ],
        "timeline": [
            "На начальном этапе frontend-модуль не имел формализованной конфигурации сборки и единых команд разработки.",
            "После инициализации npm и TypeScript были зафиксированы базовые зависимости и сценарии запуска клиентской части.",
            "Затем проведено документирование ключевых TS-компонентов через TSDoc для упрощения поддержки и ускорения вхождения новых участников.",
            "На этапе стабилизации проекта README.md был дополнен практическими инструкциями по запуску и поддержке frontend.",
            "В актуальной версии frontend-модуль имеет стандартизированную сборку и механизм автоматической генерации HTML-документации через TypeDoc.",
        ],
        "conclusion": [
            "В рамках Части 2 для frontend-модуля реализованы обязательные инженерные практики: инициализирована сборка npm/TypeScript, внедрено TSDoc-документирование и добавлена генерация HTML-документации.",
            "Цель этапа по роли Frontend UI/Electron-разработчика достигнута.",
        ],
    },
    {
        "filename": "Отчет_Часть2_Алексеев_МВ.docx",
        "developer_name": "Алексеев Максим Владимирович",
        "role": "Разработчик БД, REST API и Web-клиента",
        "intro": [
            "Цель Части 2 по данной роли заключалась в развитии серверной части для работы с базой данных и REST API, а также в подготовке сборочной инфраструктуры web-клиента.",
            "Дополнительной задачей являлось документирование Java-кода и подтверждение генерации HTML-документации.",
        ],
        "main": [
            "В backend/build.gradle добавлены зависимости PostgreSQL и Spring Data JPA для корректной интеграции с реляционной базой данных и ORM-слоем.",
            "Для DTO и REST-контроллеров авторизации подготовлены Javadoc-комментарии, описывающие назначение классов, параметры и возвращаемые значения.",
            "В директории frontend-web инициализирована система сборки на npm/package.json для web-версии клиента.",
            "Проведена тестовая проверка генерации backend-документации командой ./gradlew javadoc с успешным формированием HTML-артефактов.",
        ],
        "commits": [
            ["[ДД.ММ.2026]", "[HASH_1]", "Добавлены зависимости PostgreSQL и Spring Data JPA в backend/build.gradle"],
            ["[ДД.ММ.2026]", "[HASH_2]", "Добавлены подробные Javadoc-комментарии для DTO и REST-контроллеров авторизации"],
            ["[ДД.ММ.2026]", "[HASH_3]", "Инициализирована сборка frontend-web (package.json) и выполнена тестовая генерация ./gradlew javadoc"],
        ],
        "figures": [
            {
                "ref": "Изменения сборки backend для работы с БД представлены на рисунке 1.",
                "placeholder": "Файл backend/build.gradle: зависимости spring-boot-starter-data-jpa и postgresql",
                "caption": "Рисунок 1 - Подключение зависимостей БД в backend",
            },
            {
                "ref": "Пример Javadoc для REST API и DTO приведен на рисунке 2.",
                "placeholder": "Фрагмент кода DTO или AuthController с Javadoc-комментариями",
                "caption": "Рисунок 2 - Документирование API-слоя на Java",
            },
            {
                "ref": "Инициализация web-клиента отражена на рисунке 3.",
                "placeholder": "Файл frontend-web/package.json после инициализации npm",
                "caption": "Рисунок 3 - Конфигурация сборки web-клиента",
            },
            {
                "ref": "Результат генерации документации в консоли показан на рисунке 4.",
                "placeholder": "Консоль с выполнением ./gradlew javadoc и строкой BUILD SUCCESSFUL",
                "caption": "Рисунок 4 - Успешная генерация Javadoc",
            },
            {
                "ref": "Сгенерированная HTML-документация представлена на рисунке 5.",
                "placeholder": "Открытый в браузере файл backend/build/docs/javadoc/index.html",
                "caption": "Рисунок 5 - HTML-документация backend-модуля",
            },
            {
                "ref": "История коммитов разработчика приведена на рисунке 6.",
                "placeholder": "Вывод git log --author=\"Алексеев|Maksim\" -n 3 --oneline",
                "caption": "Рисунок 6 - История коммитов разработчика 3",
            },
        ],
        "timeline": [
            "На начальном этапе backend-часть требовала полноценной настройки доступа к реляционной базе данных и формализации API-слоя.",
            "После подключения PostgreSQL и Spring Data JPA была обеспечена базовая инфраструктура хранения и обработки данных.",
            "Затем выполнено документирование DTO и REST-контроллеров, что закрепило контракт API в понятном и проверяемом виде.",
            "Параллельно для frontend-web была инициализирована npm-сборка, чтобы унифицировать процессы запуска клиентских модулей.",
            "В актуальной версии роль БД и API закрыта рабочей конфигурацией backend и подтвержденной генерацией HTML-документации через Gradle.",
        ],
        "conclusion": [
            "По результатам Части 2 реализованы задачи по направлениям БД, REST API и web-клиента: настроены backend-зависимости, выполнено Javadoc-документирование API-слоя и подготовлена сборочная база frontend-web.",
            "Требования этапа полностью выполнены, включая подготовку артефактов для проверки (коммиты, скриншоты, HTML-документация).",
        ],
    },
]


for report in reports:
    created_path = build_report(report)
    print(created_path)
