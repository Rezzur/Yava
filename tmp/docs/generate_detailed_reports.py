# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(r"C:/Yavimax3")
OUT_DOC = ROOT / "output" / "doc"
SS = ROOT / "output" / "screenshots"

OUT_DOC.mkdir(parents=True, exist_ok=True)


def set_run(run, *, bold=False, font_name="Times New Roman", size=14):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def set_p(p, *, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, indent=True):
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)


def p_add(doc, text, *, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, bold=False, indent=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=bold)
    set_p(p, align=align, indent=indent)


def code_add(doc, lines):
    p = doc.add_paragraph()
    set_p(p, align=WD_PARAGRAPH_ALIGNMENT.LEFT, indent=False)
    for i, line in enumerate(lines):
        r = p.add_run(line)
        if i < len(lines) - 1:
            r.add_break()
        set_run(r, font_name="Courier New", size=12)


def blank(doc):
    p = doc.add_paragraph()
    set_p(p, align=WD_PARAGRAPH_ALIGNMENT.LEFT, indent=False)


def h_add(doc, text):
    p_add(doc, text, align=WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True, indent=False)


def configure_doc(doc):
    sec = doc.sections[0]
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)


def add_title_page(doc, person):
    org_lines = [
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "ФГБОУ ВО «МИРЭА - Российский технологический университет»",
        "Институт информационных технологий",
        "Кафедра математического обеспечения и стандартизации информационных технологий",
    ]
    for line in org_lines:
        p_add(doc, line, align=WD_PARAGRAPH_ALIGNMENT.CENTER, indent=False)

    blank(doc)
    blank(doc)

    p_add(doc, "ОТЧЕТ", align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, indent=False)
    p_add(doc, "по выполнению итогового проекта (части 1-4)", align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True, indent=False)
    p_add(doc, "по дисциплине «Технологии разработки программных приложений»", align=WD_PARAGRAPH_ALIGNMENT.CENTER, indent=False)
    blank(doc)
    p_add(doc, "Тема: «Кроссплатформенный десктоп-мессенджер Yavimax»", align=WD_PARAGRAPH_ALIGNMENT.CENTER, indent=False)
    blank(doc)

    p_add(doc, "Выполнил: студент группы [УКАЗАТЬ ГРУППУ]", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, indent=False)
    p_add(doc, person["full_name"], align=WD_PARAGRAPH_ALIGNMENT.RIGHT, indent=False)
    p_add(doc, f"Роль в команде: {person['role']}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, indent=False)
    p_add(doc, "Проверил: [ФИО преподавателя]", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, indent=False)

    for _ in range(9):
        blank(doc)

    p_add(doc, "Москва 2026", align=WD_PARAGRAPH_ALIGNMENT.CENTER, indent=False)
    doc.add_page_break()


def add_commit_table(doc, commits):
    table = doc.add_table(rows=1 + len(commits), cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Дата"
    table.cell(0, 1).text = "Хэш коммита"
    table.cell(0, 2).text = "Описание"

    for i, row in enumerate(commits, start=1):
        table.cell(i, 0).text = row[0]
        table.cell(i, 1).text = row[1]
        table.cell(i, 2).text = row[2]

    for tr in table.rows:
        for cell in tr.cells:
            for p in cell.paragraphs:
                set_p(p, align=WD_PARAGRAPH_ALIGNMENT.CENTER, indent=False)
                if not p.runs:
                    p.add_run("")
                for r in p.runs:
                    set_run(r)

    for i in range(1, 1 + len(commits)):
        set_p(table.cell(i, 2).paragraphs[0], align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, indent=False)


COMMON_INTRO = [
    "Итоговый проект выполнен командой из трех человек и посвящен разработке кроссплатформенного десктоп-мессенджера Yavimax. Актуальный технологический стек проекта: backend на Java/Spring Boot, клиентские части на TypeScript, сборка через Gradle и npm.",
    "Цель отчета: показать непрерывный путь разработки от первичной постановки задачи до готового к защите результата, включая сборку, документирование кода, контейнеризацию и подготовку презентационных материалов.",
    "В отчете фиксируются только реальные артефакты проекта: текущая структура репозитория, фактические конфигурационные файлы, результаты команд сборки и генерации документации, а также журнал git-коммитов.",
]

WORKFLOW_FROM_ZERO = [
    "На стартовом этапе проект существовал как идея и набор пользовательских требований: требовалось реализовать обмен сообщениями в реальном времени, базовые сценарии авторизации, чат-интерфейс и интеграцию с БД.",
    "Далее была организована репозиторная структура из трех модулей (`backend`, `frontend`, `frontend-web`) и распределены роли участников команды с привязкой к зонам ответственности.",
    "После этого выполнена техническая настройка: добавлены системы сборки Gradle и npm, описаны зависимости, подготовлены команды запуска и сборки, написаны базовые инструкции в README и RUN_INSTRUCTIONS.",
    "Следующий этап - повышение инженерной зрелости проекта: добавление Javadoc/TSDoc, подключение генерации HTML-документации, проверка воспроизводимости сборки по всем модулям.",
    "Текущая актуальная версия включает Docker-сценарий локального развертывания (backend + PostgreSQL + Redis), проверенный запуск сервисов и полный комплект материалов к защите (отчеты, презентация, скриншоты).",
]

PART1_TEXT = [
    "В части 1 были зафиксированы предметная область, целевая аудитория и базовые сценарии использования мессенджера. Ключевой задачей было согласовать единое понимание функциональности между участниками проекта до начала активной разработки.",
    "Организационно был сформирован командный процесс: выделены роли Team Lead/Backend, Frontend UI/Electron и разработчик БД/REST API/web-клиента. Такое распределение позволило параллельно развивать backend, пользовательский интерфейс и web-направление без конфликтов в ответственности.",
    "Также был определен целевой стек и структура репозитория. Это решение важно для части 2 и части 3, поскольку сборка и контейнеризация опираются на стабильную модульную архитектуру.",
    "Итогом части 1 стал переход от концепта к рабочему инженерному плану: команда получила согласованную декомпозицию задач, а проект - основу для дальнейшей реализации и демонстрации на защите.",
]

PART2_TEXT = [
    "В части 2 выполнялась инженерная подготовка к воспроизводимой разработке: настройка сборочных систем, описание зависимостей, оформление документации и фиксация команд запуска.",
    "Для backend применен Gradle. В `build.gradle` подключены зависимости Spring Boot, WebSocket/STOMP, Spring Data JPA, PostgreSQL, Redis, security-стек и OpenAPI/Swagger. Это обеспечивает серверную бизнес-логику, real-time обмен сообщениями и интеграцию с инфраструктурными сервисами.",
    "Для frontend и frontend-web использован npm с `package.json`, где описаны скрипты `build`, `dev` и генерация документации. Отдельно добавлен `typedoc` для автоматической сборки HTML-документации по TypeScript-коду.",
    "В рамках стандартизации кода добавлены Javadoc и TSDoc комментарии к ключевым компонентам. Комментарии оформлены так, чтобы на защите можно было показать не только код, но и его интерфейсный контракт (назначение, параметры, возвращаемые значения).",
    "Проверка части 2 выполнена через фактические команды: `./gradlew.bat clean test`, `./gradlew.bat javadoc`, `npm run build`, `npm run docs`, `npm run build` в `frontend-web`. По логам получены успешные сборки и сформирована HTML-документация.",
]

PART2_CODE = [
    "// backend/build.gradle (фрагмент)",
    "implementation 'org.springframework.boot:spring-boot-starter-websocket'",
    "implementation 'org.springframework.boot:spring-boot-starter-data-jpa'",
    "runtimeOnly 'org.postgresql:postgresql'",
    "javadoc { options.encoding = 'UTF-8'; options.charSet = 'UTF-8' }",
    "",
    "// frontend/package.json (фрагмент)",
    '"build": "tsc && vite build",',
    '"docs": "typedoc --entryPointStrategy expand src"',
]

PART3_TEXT = [
    "В части 3 выполнена контейнеризация backend-контура для воспроизводимого запуска проекта на любой машине с Docker Desktop. Базовая цель этапа - убрать ручные шаги разворачивания БД и кэша, а также зафиксировать единый сценарий запуска команды.",
    "Создан multi-stage `backend/Dockerfile`: на стадии `builder` приложение собирается Gradle (`bootJar`), на стадии `runtime` используется легковесный JRE-образ `eclipse-temurin:21-jre-alpine`. Такое разделение уменьшает итоговый размер runtime-образа и исключает лишние build-инструменты из прод-среды.",
    "Расширен `backend/docker-compose.yml`, где описаны три сервиса: `backend` (REST/WebSocket API), `postgres` (основное хранилище данных), `redis` (кэш/служебные сценарии обмена). Для всех сервисов заданы healthcheck-проверки, а для backend - зависимости по готовности БД и Redis.",
    "Конфигурация backend переведена на env-переменные (`DB_HOST`, `DB_PORT`, `DB_NAME`, `DB_USER`, `DB_PASSWORD`, `REDIS_HOST`, `REDIS_PORT`, `JWT_SECRET`, `SERVER_PORT`) через `application.yml` с локальными значениями по умолчанию. Это позволяет использовать один и тот же код и для Docker-режима, и для локальной разработки.",
    "Проверка части 3 выполнена командами `docker compose up -d --build` и `docker compose ps`. По журналу сервисы `yavimax_backend`, `yavimax_db`, `yavimax_redis` перешли в статус `healthy`, что является доказательством корректного контейнерного контура.",
]

PART3_CODE = [
    "# backend/Dockerfile (фрагмент)",
    "FROM gradle:9.3.1-jdk21-alpine AS builder",
    "RUN gradle --no-daemon bootJar",
    "FROM eclipse-temurin:21-jre-alpine AS runtime",
    "ENTRYPOINT [\"sh\", \"-c\", \"java ${JAVA_OPTS} -jar app.jar\"]",
    "",
    "# backend/docker-compose.yml (фрагмент)",
    "services: backend, postgres, redis",
    "depends_on: postgres/redis с condition: service_healthy",
    "healthcheck backend: wget -qO- http://localhost:8080/swagger-ui.html",
]

PART4_TEXT = [
    "В части 4 собран комплект защиты проекта: презентация на 12 слайдов, PDF-версия презентации и набор скриншотов, подтверждающих выполнение частей 1-3.",
    "Структура выступления выстроена в академической логике: постановка задачи, архитектура, реализация, сборка/документация, контейнеризация, вклад участников и итоговые выводы. Такой порядок позволяет уложиться в лимит времени и последовательно показать инженерные решения команды.",
    "Для защиты подготовлены демонстрационные материалы: экран входа, основной интерфейс чатов, HTML-документация Javadoc/TypeDoc, логи сборки и Docker-статусы. Это формирует связанный рассказ «код - сборка - запуск - результат».",
    "Дополнительно проведена верификация презентационных файлов (количество слайдов и страниц в PDF), чтобы исключить организационные риски во время защиты.",
]

CONCLUSION = [
    "Командой последовательно выполнены требования частей 1-4: от формализации задачи и выбора стека до контейнеризации и подготовки материалов к защите.",
    "Проект доведен до актуального состояния с воспроизводимой сборкой, автоматической генерацией документации и единым Docker-сценарием запуска backend-инфраструктуры.",
    "Индивидуальный вклад каждого участника зафиксирован в отчете и подтвержден таблицей коммитов и иллюстративными материалами.",
]

PEOPLE = [
    {
        "full_name": "Зырянов Ян Александрович",
        "short": "Zyryanov_YA",
        "role": "Team Lead, Backend-разработчик",
        "goal": "Цель части 2 для данной роли: обеспечить стабильную backend-сборку на Gradle, оформить серверную документацию Javadoc и подготовить инфраструктурную базу для контейнерного запуска.",
        "personal_work": [
            "Работа началась с выстраивания серверного каркаса проекта «с нуля»: определена структура backend-модуля, выбран набор зависимостей и зафиксированы принципы взаимодействия REST + WebSocket в едином приложении.",
            "На уровне сборки сформирована полноценная Gradle-конфигурация: подключены web, validation, data-jpa, redis, websocket, security и OpenAPI-зависимости, благодаря чему backend перестал быть «черновым» и стал воспроизводимо собираться на любой машине команды.",
            "С нуля настроены серверные точки запуска и разработческий контур: оформлены README и RUN_INSTRUCTIONS с единым сценарием команд, чтобы любой участник мог быстро поднять backend без ручного восстановления контекста.",
            "Отдельно выполнена инженерная проработка документации: добавлены Javadoc-комментарии к ключевым конфигурационным классам, а генерация Javadoc приведена к корректной UTF-8-выгрузке, что критично для русскоязычных пояснений в учебном проекте.",
            "В части real-time логики была выстроена WebSocket-конфигурация STOMP (префиксы брокера, пользовательские очереди, авторизация по JWT), что обеспечило корректную маршрутизацию событий между сервером и клиентами.",
            "На этапе контейнеризации backend был переведен в multi-stage Docker-процесс: сборка `bootJar` на builder-слое и отдельный runtime-слой для запуска, что уменьшило размер итогового образа и улучшило повторяемость запуска.",
            "Дополнительно реализован compose-контур из трех сервисов (backend + PostgreSQL + Redis) с healthcheck-проверками, зависимостями по готовности и переменными окружения; это позволило стабильно запускать весь серверный стек одной командой.",
            "В результате личного вклада backend прошел путь от базовой заготовки до актуальной рабочей версии с проверенной сборкой, документацией и контейнерным запуском, готовым к демонстрации на защите.",
        ],
        "personal_snippets": [
            {
                "title": "Конфигурация STOMP-брокера и user-очередей в WebSocketConfig",
                "context": "Фрагмент показывает, как на backend заданы каналы маршрутизации для чатов и персональных событий пользователей.",
                "code": [
                    "@Override",
                    "public void configureMessageBroker(MessageBrokerRegistry config) {",
                    "    config.enableSimpleBroker(\"/queue\", \"/topic\");",
                    "    config.setApplicationDestinationPrefixes(\"/app\");",
                    "    config.setUserDestinationPrefix(\"/user\");",
                    "}",
                ],
            },
            {
                "title": "Регистрация защищенного WebSocket endpoint с SockJS fallback",
                "context": "Этот фрагмент отражает практическую часть real-time контура: endpoint `/ws` и подключение handshake-interceptor для токенной авторизации.",
                "code": [
                    "WebSocketHandshakeInterceptor authInterceptor = new WebSocketHandshakeInterceptor(jwtService);",
                    "",
                    "registry.addEndpoint(\"/ws\")",
                    "        .setAllowedOriginPatterns(\"*\")",
                    "        .addInterceptors(authInterceptor)",
                    "        .withSockJS();",
                ],
            },
            {
                "title": "Docker runtime-процесс backend",
                "context": "Фрагмент демонстрирует переход от локального запуска к воспроизводимому контейнерному сценарию с разделением build/runtime.",
                "code": [
                    "FROM gradle:9.3.1-jdk21-alpine AS builder",
                    "RUN gradle --no-daemon bootJar",
                    "",
                    "FROM eclipse-temurin:21-jre-alpine AS runtime",
                    "COPY --from=builder /app/build/libs/*.jar ./app.jar",
                    "ENTRYPOINT [\"sh\", \"-c\", \"java ${JAVA_OPTS} -jar app.jar\"]",
                ],
            },
        ],
        "commits": [
            ("2026-04-02", "3fd57ce", "docs: implement README, Javadoc and configure documentation build systems"),
            ("2026-03-22", "8b8d6ba", "feat: implement real-time ws broadcast, mark-as-read, and initial-based avatars"),
            ("2026-03-22", "1099b48", "fix: resolve compilation errors in GlobalExceptionHandler"),
        ],
        "figures": [
            ("Структура репозитория, сформированная на этапе части 1, представлена на рисунке {n}.", "00_project_structure.png", "Рисунок {n} - Структура репозитория Yavimax"),
            ("Состав backend-зависимостей и настройки Gradle приведены на рисунке {n}.", "01_build_gradle.png", "Рисунок {n} - Фрагмент файла backend/build.gradle"),
            ("Персональный фрагмент backend-кода участника показан на рисунке {n}.", "24_zyryanov_personal_code.png", "Рисунок {n} - Пример кода Зырянова Я.А."),
            ("Пример Javadoc-комментариев для backend-конфигурации показан на рисунке {n}.", "04_websocket_javadoc.png", "Рисунок {n} - Javadoc в WebSocketConfig.java"),
            ("Результат проверки backend-сборки командой `clean test` показан на рисунке {n}.", "06_console_backend_test.png", "Рисунок {n} - Успешная backend-сборка (Gradle test)"),
            ("Успешная генерация Javadoc отражена на рисунке {n}.", "07_console_backend_javadoc.png", "Рисунок {n} - Лог генерации Javadoc"),
            ("Содержимое Dockerfile backend-контейнера представлено на рисунке {n}.", "19_backend_dockerfile.png", "Рисунок {n} - Файл backend/Dockerfile"),
            ("Конфигурация docker-compose для трех сервисов приведена на рисунке {n}.", "20_backend_docker_compose.png", "Рисунок {n} - Файл backend/docker-compose.yml"),
            ("Переход на env-конфигурацию backend показан на рисунке {n}.", "21_backend_application_yml.png", "Рисунок {n} - Файл backend/src/main/resources/application.yml"),
            ("Лог запуска контейнерного контура показан на рисунке {n}.", "11_console_docker_up.png", "Рисунок {n} - Выполнение docker compose up -d --build"),
            ("Подтверждение готовности сервисов в Docker показано на рисунке {n}.", "12_console_docker_ps.png", "Рисунок {n} - Статусы контейнеров (healthy)"),
            ("Сформированная HTML-документация backend приведена на рисунке {n}.", "14_javadoc_browser.png", "Рисунок {n} - Javadoc в браузере"),
            ("Фрагмент реального журнала git-коммитов показан на рисунке {n}.", "13_console_git_log.png", "Рисунок {n} - Фрагмент git log"),
            ("Проверка итоговых файлов презентации показана на рисунке {n}.", "23_presentation_files.png", "Рисунок {n} - Сформированные файлы презентации"),
            ("Верификация 12 слайдов презентации и PDF-версии представлена на рисунке {n}.", "22_presentation_validation.png", "Рисунок {n} - Проверка презентационных артефактов"),
            ("Экран входа в приложение для демонстрации на защите приведен на рисунке {n}.", "17_ui_login.png", "Рисунок {n} - Экран авторизации приложения"),
            ("Основной экран мессенджера, используемый в демо-сценарии, показан на рисунке {n}.", "18_ui_main.png", "Рисунок {n} - Основной экран чатов"),
        ],
    },
    {
        "full_name": "Топал Виктор Викторович",
        "short": "Topal_VV",
        "role": "Frontend UI/Electron-разработчик",
        "goal": "Цель части 2 для данной роли: настроить frontend-сборку на npm/TypeScript, документировать клиентский код в TSDoc и обеспечить генерацию TypeDoc-документации.",
        "personal_work": [
            "Работа по клиентской части началась с нуля: был сформирован каркас фронтенд-модуля, выстроены базовые слои UI, API и state-management, после чего проект получил стабильную основу для быстрого расширения функционала.",
            "На сборочном уровне настроен npm/TypeScript pipeline (`dev`, `build`, `preview`, `docs`), что обеспечило предсказуемую разработку интерфейса и воспроизводимую production-сборку без ручных шагов.",
            "Существенная часть вклада связана с интеграцией frontend и backend: реализованы HTTP-вызовы через единый API-клиент, обработка токена и 401-сценариев, а также синхронизация состояния после авторизации.",
            "Для realtime-поведения настроено WebSocket-взаимодействие: подключение STOMP/SockJS, подписка на персональные и чатовые каналы, доставка входящих сообщений в store и обновление интерфейса без перезагрузки страницы.",
            "С нуля развивалась логика интерфейса чата: отправка текстовых сообщений, обработка вложений, работа с голосовыми сообщениями, автоматическое обновление последнего сообщения в сайдбаре и счетчиков непрочитанных.",
            "Отдельное внимание уделено архитектуре клиентского состояния: через Zustand реализованы сущности `chats`, `messages`, `activeChat`, сценарии загрузки истории и mark-as-read, что дало управляемую и читаемую бизнес-логику на фронтенде.",
            "В части документации добавлены TSDoc-комментарии для ключевых модулей (`client.ts`, `websocket.ts`, `useUser.ts`) и включена генерация TypeDoc, чтобы на защите можно было показать не только интерфейс, но и инженерную прозрачность кода.",
            "Итог личного вклада: frontend прошел путь от начальной заготовки до актуального интерактивного интерфейса, связанного с backend в реальном времени и подтвержденного сборкой и документацией.",
        ],
        "personal_snippets": [
            {
                "title": "Единый API-клиент с автоматической авторизацией",
                "context": "Фрагмент показывает, как на клиенте централизована работа с токеном и обработка невалидной сессии.",
                "code": [
                    "export const api = axios.create({ baseURL: API_BASE_URL });",
                    "",
                    "api.interceptors.request.use((config) => {",
                    "  const token = localStorage.getItem('yavimax_token');",
                    "  if (token) config.headers.Authorization = `Bearer ${token}`;",
                    "  return config;",
                    "});",
                ],
            },
            {
                "title": "WebSocket-подключение и подписка на персональный канал",
                "context": "Фрагмент демонстрирует realtime-интеграцию клиента: подключение к STOMP и обработку входящих событий `message.created`.",
                "code": [
                    "const client = new Client({",
                    "  webSocketFactory: () => new SockJS(`${SOCKET_URL}?token=${token}`),",
                    "  connectHeaders: { Authorization: `Bearer ${token}` },",
                    "  reconnectDelay: 5000,",
                    "});",
                    "",
                    "client.subscribe('/user/queue/messages', (message) => {",
                    "  const data = JSON.parse(message.body);",
                    "  if (data.type === 'message.created') addMessage(data.payload);",
                    "});",
                ],
            },
            {
                "title": "Логика store для обновления чатов и unread-счетчиков",
                "context": "Фрагмент иллюстрирует работу клиентского состояния: актуализацию сайдбара и корректную обработку активного/неактивного чата.",
                "code": [
                    "addMessage: (message) => {",
                    "  const state = get();",
                    "  get().updateChatLastMessage(message.chatId, message.text, message.timestamp);",
                    "  if (state.activeChatId !== message.chatId) {",
                    "    set({ chats: state.chats.map(chat =>",
                    "      chat.id === message.chatId ? { ...chat, unreadCount: (chat.unreadCount || 0) + 1 } : chat",
                    "    )});",
                    "    return;",
                    "  }",
                    "},",
                ],
            },
        ],
        "commits": [
            ("2026-03-21", "5d5eef0", "feat: connect frontend and backend for search and real-time messaging"),
            ("2026-04-02", "58b2758", "docs: generate TypeDoc and fix frontend TS errors"),
            ("2026-04-02", "c60b725", "chore: cleanup unused UI features and fix media upload bugs"),
        ],
        "figures": [
            ("Структура репозитория и выделение frontend-модуля представлены на рисунке {n}.", "00_project_structure.png", "Рисунок {n} - Структура репозитория Yavimax"),
            ("Сборочная конфигурация frontend-модуля приведена на рисунке {n}.", "02_frontend_package_json.png", "Рисунок {n} - Файл frontend/package.json"),
            ("Персональный фрагмент frontend-кода участника приведен на рисунке {n}.", "25_topal_personal_code.png", "Рисунок {n} - Пример кода Топала В.В."),
            ("Пример TSDoc-комментариев в клиентском API показан на рисунке {n}.", "05_frontend_tsdoc_client.png", "Рисунок {n} - TSDoc в frontend/src/shared/api/client.ts"),
            ("Результат production-сборки frontend приведен на рисунке {n}.", "08_console_frontend_build.png", "Рисунок {n} - Лог npm run build (frontend)"),
            ("Результат генерации TypeDoc показан на рисунке {n}.", "09_console_frontend_docs.png", "Рисунок {n} - Лог npm run docs (frontend)"),
            ("Сформированная HTML-документация frontend показана на рисунке {n}.", "15_typedoc_browser.png", "Рисунок {n} - TypeDoc в браузере"),
            ("Контейнерный запуск серверной части для совместной проверки UI приведен на рисунке {n}.", "11_console_docker_up.png", "Рисунок {n} - Docker-запуск backend-контура"),
            ("Подтверждение готовности backend/DB/Redis в Docker показано на рисунке {n}.", "12_console_docker_ps.png", "Рисунок {n} - Статусы контейнеров (healthy)"),
            ("Фрагмент реального журнала коммитов команды приведен на рисунке {n}.", "13_console_git_log.png", "Рисунок {n} - Фрагмент git log"),
            ("Проверка сформированных файлов презентации показана на рисунке {n}.", "23_presentation_files.png", "Рисунок {n} - Файлы презентации для защиты"),
            ("Проверка количества слайдов и страниц презентации приведена на рисунке {n}.", "22_presentation_validation.png", "Рисунок {n} - Верификация презентации"),
            ("Экран авторизации, используемый в демонстрации UX, показан на рисунке {n}.", "17_ui_login.png", "Рисунок {n} - Экран входа"),
            ("Основной экран мессенджера с чатом показан на рисунке {n}.", "18_ui_main.png", "Рисунок {n} - Основной экран интерфейса"),
        ],
    },
    {
        "full_name": "Алексеев Максим Владимирович",
        "short": "Alekseev_MV",
        "role": "Разработчик БД, REST API и web-клиента",
        "goal": "Цель части 2 для данной роли: поддержать backend-сборку с зависимостями БД/REST, обеспечить документирование серверных API и стабилизировать pipeline `frontend-web`.",
        "personal_work": [
            "Разработка в зоне ответственности «БД + REST API + web-клиент» также велась фактически с нуля: требовалось связать серверную бизнес-логику, структуру хранения данных и клиентские сценарии обращения к API в единый рабочий контур.",
            "На backend-уровне обеспечена корректная интеграция с PostgreSQL через Spring Data JPA: проверены зависимости, репозитории и сущности, после чего серверная часть получила устойчивую основу для хранения пользователей, чатов и сообщений.",
            "Значительная часть вклада была направлена на стабилизацию REST-слоя: доработаны и выровнены DTO/контроллеры, устранены сбои в серверной обработке ошибок и приведены к рабочему состоянию ключевые endpoint-сценарии авторизации и обмена сообщениями.",
            "В service-слое прорабатывалась логика создания приватных чатов, отправки сообщений, обновления unread-статусов и трансляции событий клиентам, что обеспечило связность между БД, REST и WebSocket-механикой.",
            "Параллельно поддерживалась интеграция backend с UI: серверные ответы и формат DTO синхронизировались с потребностями фронтенда, чтобы убрать рассинхрон по полям и обеспечить корректную отрисовку клиентской части.",
            "По направлению `frontend-web` сформирован и проверен самостоятельный build pipeline, что позволило получить третью сборочную единицу проекта с воспроизводимой командой `npm run build`.",
            "На этапе контейнеризации внесен вклад в проверку инфраструктурного запуска: контроль корректности compose-связок, проверка готовности сервисов и валидация доступности API после старта окружения.",
            "Итог личного вклада: модуль БД/REST/web прошел путь от разрозненных заготовок до согласованной и проверенной части общей системы, устойчиво работающей в локальном и контейнерном режиме.",
        ],
        "personal_snippets": [
            {
                "title": "Сервис авторизации: регистрация и выдача JWT",
                "context": "Фрагмент показывает серверную логику валидации уникальности пользователя и формирования токена после успешной регистрации.",
                "code": [
                    "@Transactional",
                    "public AuthResponse register(RegisterRequest request) {",
                    "    if (userRepository.findByUsername(request.getUsername()).isPresent()) {",
                    "        throw new AuthException(\"Username already exists\");",
                    "    }",
                    "    User savedUser = userRepository.save(user);",
                    "    String token = jwtService.generateToken(new UserDetailsImpl(savedUser));",
                    "    return buildAuthResponse(savedUser, token);",
                    "}",
                ],
            },
            {
                "title": "Сервис чатов: отправка сообщения и публикация события",
                "context": "Ключевой фрагмент бизнес-логики: сохранение сообщения в БД и рассылка события в чатовый и персональные каналы.",
                "code": [
                    "Message saved = messageRepository.save(message);",
                    "MessageDto dto = toMessageDto(saved);",
                    "ChatMessageHandler.WebSocketMessage<MessageDto> wsMessage = ...;",
                    "messagingTemplate.convertAndSend(\"/topic/chats.\" + chatId, wsMessage);",
                    "for (ChatMember member : members) {",
                    "  messagingTemplate.convertAndSendToUser(member.getUser().getUsername(), \"/queue/messages\", wsMessage);",
                    "}",
                ],
            },
            {
                "title": "Сервис чатов: расчет данных для сайдбара (last message/unread)",
                "context": "Фрагмент демонстрирует, как backend готовит агрегированные данные для корректной работы клиентского списка чатов.",
                "code": [
                    "messageRepository.findFirstByChatIdOrderByCreatedAtDesc(chat.getId()).ifPresent(msg -> {",
                    "    builder.lastMessage(msg.getText() != null ? msg.getText() : \"\");",
                    "    builder.timestamp(formatTimestamp(msg.getCreatedAt()));",
                    "});",
                    "",
                    "Long unreadCount = messageRepository.countUnreadMessages(chat.getId(), currentUserId);",
                    "builder.unreadCount(unreadCount != null ? unreadCount.intValue() : 0);",
                ],
            },
        ],
        "commits": [
            ("2026-03-21", "df95033", "Initial commit: Fix build issues in backend and frontend"),
            ("2026-04-02", "80fab0f", "fix: implement media uploads, voice messages and websocket synchronization"),
            ("2026-03-25", "7043ecf", "fix: resolve ws sync issues, message duplication, and improve avatar UI"),
        ],
        "figures": [
            ("Структура репозитория и разделение на backend/frontend/frontend-web показаны на рисунке {n}.", "00_project_structure.png", "Рисунок {n} - Структура репозитория Yavimax"),
            ("Конфигурация backend-сборки с зависимостями БД приведена на рисунке {n}.", "01_build_gradle.png", "Рисунок {n} - Файл backend/build.gradle"),
            ("Персональный фрагмент серверного кода участника показан на рисунке {n}.", "26_alekseev_personal_code.png", "Рисунок {n} - Пример кода Алексеева М.В."),
            ("Сборочная конфигурация web-клиента показана на рисунке {n}.", "03_frontendweb_package_json.png", "Рисунок {n} - Файл frontend-web/package.json"),
            ("Результат сборки frontend-web приведен на рисунке {n}.", "10_console_frontendweb_build.png", "Рисунок {n} - Лог npm run build (frontend-web)"),
            ("Пример Javadoc-комментариев backend-конфигурации показан на рисунке {n}.", "04_websocket_javadoc.png", "Рисунок {n} - Javadoc в backend-коде"),
            ("Сформированная backend HTML-документация показана на рисунке {n}.", "14_javadoc_browser.png", "Рисунок {n} - Javadoc в браузере"),
            ("Сформированная frontend HTML-документация показана на рисунке {n}.", "15_typedoc_browser.png", "Рисунок {n} - TypeDoc в браузере"),
            ("Содержимое Dockerfile для backend-сервиса приведено на рисунке {n}.", "19_backend_dockerfile.png", "Рисунок {n} - Файл backend/Dockerfile"),
            ("Конфигурация контейнерного контура в compose-файле показана на рисунке {n}.", "20_backend_docker_compose.png", "Рисунок {n} - Файл backend/docker-compose.yml"),
            ("Лог старта контейнеров приведен на рисунке {n}.", "11_console_docker_up.png", "Рисунок {n} - Docker compose up"),
            ("Статусы сервисов после старта приведены на рисунке {n}.", "12_console_docker_ps.png", "Рисунок {n} - Docker compose ps"),
            ("Фрагмент журнала коммитов для подтверждения вклада приведен на рисунке {n}.", "13_console_git_log.png", "Рисунок {n} - Фрагмент git log"),
            ("Проверка сформированных презентационных файлов показана на рисунке {n}.", "23_presentation_files.png", "Рисунок {n} - Файлы презентации"),
            ("Проверка количества слайдов/страниц презентации приведена на рисунке {n}.", "22_presentation_validation.png", "Рисунок {n} - Верификация презентации"),
            ("Экран входа для демонстрации итогового продукта показан на рисунке {n}.", "17_ui_login.png", "Рисунок {n} - Экран авторизации"),
            ("Главный экран чатов для демо-сценария показан на рисунке {n}.", "18_ui_main.png", "Рисунок {n} - Основной экран мессенджера"),
        ],
    },
]


def add_common_sections(doc, person):
    h_add(doc, "Введение")
    for t in COMMON_INTRO:
        p_add(doc, t)
    p_add(doc, person["goal"])

    h_add(doc, "1. Краткая проделанная работа (от 0 до актуальной версии)")
    for t in WORKFLOW_FROM_ZERO:
        p_add(doc, t)

    h_add(doc, "2. Выполнение части 1: постановка задачи и организация проекта")
    for t in PART1_TEXT:
        p_add(doc, t)

    h_add(doc, "3. Выполнение части 2: сборка и документация")
    for t in PART2_TEXT:
        p_add(doc, t)

    p_add(doc, "Ниже приведены ключевые фрагменты конфигурации, использованные в части 2.")
    code_add(doc, PART2_CODE)

    h_add(doc, "4. Выполнение части 3: контейнеризация и Docker-сценарий")
    for t in PART3_TEXT:
        p_add(doc, t)

    p_add(doc, "Ниже приведены характерные фрагменты Docker-конфигурации проекта.")
    code_add(doc, PART3_CODE)

    h_add(doc, "5. Выполнение части 4: подготовка защиты проекта")
    for t in PART4_TEXT:
        p_add(doc, t)

    h_add(doc, "6. Личный вклад участника")
    for t in person["personal_work"]:
        p_add(doc, t)

    snippets = person.get("personal_snippets", [])
    if snippets:
        h_add(doc, "6.1 Ключевые фрагменты кода участника")
        for i, snippet in enumerate(snippets, start=1):
            p_add(doc, f"Фрагмент {i}: {snippet['title']}", bold=True)
            p_add(doc, snippet["context"])
            code_add(doc, snippet["code"])

    h_add(doc, "7. Таблица коммитов")
    p_add(
        doc,
        "В проекте использовался общий командный аккаунт Git, поэтому в таблице зафиксированы реальные хэши из фактического журнала `git log`, сопоставленные с зоной ответственности участника.",
    )
    add_commit_table(doc, person["commits"])
    blank(doc)


def add_figures(doc, person):
    h_add(doc, "8. Иллюстрации и подтверждающие материалы")
    fig_no = 1
    for ref, img_name, cap_tmpl in person["figures"]:
        p_add(doc, ref.format(n=fig_no))
        img_path = SS / img_name
        if img_path.exists():
            doc.add_picture(str(img_path), width=Cm(15.8))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            p_add(doc, f"[ВСТАВИТЬ СКРИНШОТ: {img_name}]", align=WD_PARAGRAPH_ALIGNMENT.CENTER, indent=False)
        p_add(doc, cap_tmpl.format(n=fig_no), align=WD_PARAGRAPH_ALIGNMENT.CENTER, indent=False)
        fig_no += 1


def add_conclusion(doc):
    h_add(doc, "9. Заключение")
    for t in CONCLUSION:
        p_add(doc, t)


def build_reports():
    for person in PEOPLE:
        doc = Document()
        configure_doc(doc)
        add_title_page(doc, person)
        add_common_sections(doc, person)
        add_figures(doc, person)
        add_conclusion(doc)

        ascii_name = f"Final_Report_Parts1-4_{person['short']}_detailed.docx"
        out_ascii = OUT_DOC / ascii_name
        try:
            doc.save(out_ascii)
        except PermissionError:
            out_ascii = OUT_DOC / f"Final_Report_Parts1-4_{person['short']}_detailed_v2.docx"
            doc.save(out_ascii)

        ru_name_map = {
            "Zyryanov_YA": "Отчет_Итоговый_Части1-4_Зырянов_ЯА_подробный.docx",
            "Topal_VV": "Отчет_Итоговый_Части1-4_Топал_ВВ_подробный.docx",
            "Alekseev_MV": "Отчет_Итоговый_Части1-4_Алексеев_МВ_подробный.docx",
        }
        out_ru = OUT_DOC / ru_name_map[person["short"]]
        try:
            doc.save(out_ru)
        except PermissionError:
            fallback_ru = ru_name_map[person["short"]].replace(".docx", "_v2.docx")
            out_ru = OUT_DOC / fallback_ru
            doc.save(out_ru)

        print(out_ascii)
        print(out_ru)


if __name__ == "__main__":
    build_reports()
